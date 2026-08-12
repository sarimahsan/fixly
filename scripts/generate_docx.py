import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = docx.Document()

    PRIMARY_NAVY = RGBColor(15, 23, 42)      # #0F172A Dark Slate
    ACCENT_BLUE = RGBColor(37, 99, 235)      # #2563EB Royal Blue
    TEXT_DARK = RGBColor(51, 65, 85)        # #334155 Slate Text
    MUTED_GRAY = RGBColor(100, 116, 139)    # #64748B Gray Text
    WHITE = RGBColor(255, 255, 255)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def set_cell_shading(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{margin}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run("FIXLY — System Architecture & Specification")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_NAVY

    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run("Autonomous Incident Detection & Self-Healing Platform | Hackathon Proposal")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12.5)
    run_sub.font.italic = True
    run_sub.font.color.rgb = ACCENT_BLUE
    p_sub.paragraph_format.space_after = Pt(18)

    p_hr = doc.add_paragraph()
    p_hr_run = p_hr.add_run("―" * 55)
    p_hr_run.font.color.rgb = MUTED_GRAY
    p_hr.paragraph_format.space_after = Pt(18)

    def add_heading_1(text):
        h = doc.add_heading(level=1)
        r = h.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_NAVY
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(8)
        return h

    def add_heading_2(text):
        h = doc.add_heading(level=2)
        r = h.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(12.5)
        r.font.bold = True
        r.font.color.rgb = ACCENT_BLUE
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        return h

    def add_body_paragraph(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(10.5)
        r.font.color.rgb = TEXT_DARK
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        return p

    def add_bullet_point(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Arial'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = PRIMARY_NAVY
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(10.5)
        r.font.color.rgb = TEXT_DARK
        p.paragraph_format.space_after = Pt(4)
        return p

    add_heading_1("1. Executive Summary & Vision")
    add_body_paragraph(
        "Fixly is an enterprise-grade, AI-driven autonomous incident detection and self-healing platform. "
        "Designed for modern cloud applications and remote server infrastructures, Fixly continuously monitors target server environments "
        "over secure SSH, ingests live application logs and system vitals (CPU, RAM, Disk), deduplicates repeated errors in real time, "
        "diagnoses root causes using Advanced Generative AI (Groq Llama 3.3 / Claude / GPT-4o), and automatically generates, tests, and commits verified code fixes directly to Git repositories."
    )

    add_heading_1("2. Technology Stack & Specifications")
    
    table_stack = doc.add_table(rows=1, cols=3)
    table_stack.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table_stack.rows[0].cells
    headers = ["System Layer", "Technology Selection", "Technical Rationale"]
    widths = [Inches(1.8), Inches(2.3), Inches(2.4)]

    for idx, (cell, text, width) in enumerate(zip(hdr_cells, headers, widths)):
        cell.width = width
        set_cell_shading(cell, "0F172A")
        set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = WHITE

    stack_data = [
        ("Language & Runtime", "JavaScript (Node.js v20+)", "High asynchronous I/O performance, unified JavaScript stack across client & server."),
        ("Frontend Dashboard", "React 18+ with Vite & Tailwind CSS", "Utility-first responsive design, lightning-fast HMR, modular component architecture."),
        ("Backend Services API", "Express / Fastify (Node.js)", "Lightweight HTTP routing, native JSON parsing, robust middleware architecture."),
        ("Real-Time Engine", "WebSockets (ws / socket.io)", "Bi-directional low-latency event broadcasting for live incident feeds and vitals gauges."),
        ("Database & ODM Layer", "MongoDB with Mongoose ORM", "Schema validation models, index acceleration, and atomic document updates."),
        ("SSH Transport", "node-ssh / ssh2", "Encrypted key-based remote log streaming and vitals command execution."),
        ("Version Control Integration", "simple-git", "Programmatic Git repository management, automated feature branch creation, and PR pushing."),
        ("AI Integration Options", "Groq Llama 3.3 (llama-3.3-70b-versatile) / Claude-3.5 / GPT-4o", "High-speed LPU Llama 3.3 option for ultra-fast root cause diagnosis and unified diff generation (~1.4s).")
    ]

    for row_idx, (layer, tech, rationale) in enumerate(stack_data):
        row_cells = table_stack.add_row().cells
        bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for idx, (cell, text, width) in enumerate(zip(row_cells, [layer, tech, rationale], widths)):
            cell.width = width
            set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.color.rgb = PRIMARY_NAVY if idx == 0 else TEXT_DARK
            if idx == 0:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    add_heading_1("3. System Modules & Functional Architecture")
    add_heading_2("Module 1: Server Connection & Monitoring (User 1 Scope)")
    add_bullet_point("Key-Based SSH Transport (ssh_client.js): ", "Connects to remote target server via node-ssh private key authentication.")
    add_bullet_point("Continuous Log Streaming (log_reader.js): ", "Tails application logs line-by-line in real time, filtering error stack traces.")
    add_bullet_point("SHA-256 Deduplication Engine (dedup_engine.js): ", "Normalizes logs by stripping variable IDs/timestamps and generates a SHA-256 fingerprint hash.")
    add_bullet_point("Real-Time Broadcaster (ws_broadcaster.js): ", "Emits live incident:created, incident:updated, and vitals:updated events over WebSocket.")

    add_heading_2("Module 2: AI Logic & Version Control Integration (User 2 Scope)")
    add_bullet_point("AI Diagnosis & Patch Engine (diagnosis.js, code_fixer.js): ", "Analyzes error context using Groq Llama 3.3 (llama-3.3-70b-versatile) or Claude/GPT-4o to return root cause, confidence score, and unified diff patches.")
    add_bullet_point("Git Manager & PR Creator (git_client.js): ", "Creates dedicated branch fix/inc-<id> via simple-git, commits patch, and pushes Pull Request.")
    add_bullet_point("Recovery Verifier (recovery.js): ", "Monitors log stream post-fix. Auto-resolves incident when zero recurrences are verified.")

    add_heading_2("Module 3: Interface & User Experience (User 3 Scope)")
    add_bullet_point("Live Incident Feed (LiveFeed.jsx): ", "Real-time updating visual feed styled with Tailwind CSS severity badges.")
    add_bullet_point("Code Diff Inspector (DiffViewer.jsx): ", "Line-by-line visual before/after code comparison modal.")

    add_heading_2("Module 4: Auth, Access Control & System Lead (User 4 Scope)")
    add_bullet_point("JWT Auth & RBAC (auth_service.js): ", "Issues authenticated JWT tokens carrying ADMIN, OPERATOR, and READ_ONLY role claims.")
    add_bullet_point("Backend Settings API (settings_service.js): ", "Provides encrypted GET/PUT /api/settings endpoints for masked token and AI provider selection.")

    add_heading_1("4. End-to-End System Workflow")
    add_bullet_point("Step 1 (Ingestion & Dedup): ", "An unhandled exception occurs on the target server. Module 1 reads log line over SSH, calculates SHA-256 fingerprint, and updates MongoDB.")
    add_bullet_point("Step 2 (AI Patch Generation): ", "Module 2 sends stack trace to Groq Llama-3.3-70B (or Claude/GPT-4o), generates unified diff patch, and opens Git Pull Request.")
    add_bullet_point("Step 3 (Live Stream & Verification): ", "Module 3 presents live incident on React dashboard. Module 2 auto-resolves incident after zero-recurrence window.")

    output_dir = os.path.join(os.getcwd(), "docs")
    os.makedirs(output_dir, exist_ok=True)
    doc_path = os.path.join(output_dir, "Fixly_System_Specification_v3.docx")
    try:
        doc.save(doc_path)
        print(f"Document successfully saved at: {doc_path}")
    except Exception as e:
        print(f"Save error: {e}")

if __name__ == "__main__":
    create_document()
